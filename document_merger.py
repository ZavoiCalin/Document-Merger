import os
from docx import Document

def merge_word_docs_in_directory(output_file):
    # Create a new Document object to store the merged contents
    merged_doc = Document()

    # Get a list of all .docx files in the current directory
    docx_files = [file for file in os.listdir() if file.endswith(".docx")]

    # Iterate over each input document
    for file in docx_files:
        # Open the input document
        doc = Document(file)

        # Iterate over each paragraph in the input document
        for paragraph in doc.paragraphs:
            # Check if the paragraph is a heading (i.e., a chapter, subchapter, etc.)
            if paragraph.style.name.startswith("Heading"):
                # Determine the level of the heading (chapter, subchapter, sub-subchapter, etc.)
                level = int(paragraph.style.name[-1])

                # Set the style of the paragraph to the corresponding subchapter style in the merged document
                new_style = "Heading " + str(level + 1)  # Increment the level for subchapters
                merged_paragraph = merged_doc.add_paragraph(paragraph.text, style=new_style)

                # Add a prefix to the text of the merged paragraph
                prefix = "CHAPTER" + "." * level  # Add dots to represent the level of the chapter
                merged_paragraph.text = prefix + " " + merged_paragraph.text

            # If the paragraph is not a heading, simply copy it to the merged document
            else:
                merged_doc.add_paragraph(paragraph.text)

    # Save the merged document to the output file
    merged_doc.save(output_file)

if __name__ == "__main__":
    # Output file where the merged document will be saved
    output_file = "merged_document.docx"

    # Merge the .docx documents in the current directory
    merge_word_docs_in_directory(output_file)

    print("Documents merged successfully!")